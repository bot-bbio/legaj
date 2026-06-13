import os
import sys
import logging

from _encoding import force_utf8_io

# Resumes routinely contain non-Latin-1 glyphs (bullets, em-dashes, smart
# quotes, accented names) that crash on Windows' default cp1252 stdout.
force_utf8_io()

# pypdf logs "Ignoring wrong pointing object" warnings for many real-world
# PDFs; they are non-fatal and only add noise to our stderr output.
logging.getLogger("pypdf").setLevel(logging.ERROR)

def _die(msg):
    print(msg, file=sys.stderr)
    sys.exit(1)

def parse_pdf(file_path):
    try:
        from pypdf import PdfReader
    except ImportError:
        _die("Error: pypdf library is not installed. Run `pip install -r requirements.txt`.")

    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception as e:
        _die(f"Error reading PDF file: {str(e)}")

def parse_docx(file_path):
    try:
        import docx
    except ImportError:
        _die("Error: python-docx library is not installed. Run `pip install -r requirements.txt`.")

    try:
        doc = docx.Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                text += " | ".join(row_text) + "\n"
        return text.strip()
    except Exception as e:
        _die(f"Error reading DOCX file: {str(e)}")

def parse_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        _die(f"Error reading text file: {str(e)}")

def main():
    if len(sys.argv) < 2:
        _die("Usage: python parse_resume.py <resume_path>")

    file_path = sys.argv[1]
    if not os.path.exists(file_path):
        _die(f"Error: File '{file_path}' does not exist.")

    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.pdf':
        text = parse_pdf(file_path)
    elif ext == '.docx':
        text = parse_docx(file_path)
    elif ext == '.doc':
        _die("Error: Legacy .doc format is not supported. Please save the file as .docx or PDF and try again.")
    elif ext in ['.txt', '.md']:
        text = parse_txt(file_path)
    else:
        _die(f"Error: Unsupported file format '{ext}'. Supported formats: PDF, DOCX, TXT, MD.")

    if not text or not text.strip():
        _die(f"Error: No text could be extracted from '{file_path}'. The file may be empty or image-only.")

    print(text)

if __name__ == "__main__":
    main()
